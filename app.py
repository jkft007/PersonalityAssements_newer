import gradio as gr
import json
import requests
from datetime import datetime
import os
from pathlib import Path
import platform
from dotenv import load_dotenv
import random
import re
import tempfile
import uuid

# Questions and answer options
questions = [
    {
        "id": 1,
        "category": "Stress Response",
        "question": "When faced with major stress, what do you typically do?",
        "options": {
            "Narcissistic": "I push harder, prove I can handle it better than others",
            "Obsessive": "I make lists, analyze all options systematically",
            "Depressive": "I feel it's somehow my fault, withdraw to reflect",
            "Paranoid": "I suspect others are working against me",
            "Schizoid": "I retreat into solitude, need space from people",
            "Hysterical": "I become emotional, seek immediate support",
            "Borderline": "My feelings swing wildly, I might act impulsively",
            "Masochistic": "I endure it, others have it worse"
        }
    },
    {
        "id": 2,
        "category": "Self-Worth Source",
        "question": "What makes you feel valuable?",
        "options": {
            "Narcissistic": "Achievement, admiration, being exceptional",
            "Obsessive": "Doing things correctly, being reliable",
            "Depressive": "Helping others, being needed",
            "Paranoid": "Being vigilant, protecting myself/others",
            "Schizoid": "My inner world, autonomy",
            "Hysterical": "Being attractive, emotionally connecting",
            "Borderline": "Varies drastically day-to-day",
            "Masochistic": "Sacrifice, suffering for others"
        }
    },
    {
        "id": 3,
        "category": "Relationship Pattern",
        "question": "Describe your typical relationship dynamic:",
        "options": {
            "Narcissistic": "I need admiration; partners disappoint me",
            "Obsessive": "I'm loyal but critical of imperfection",
            "Depressive": "I give endlessly, fear abandonment",
            "Paranoid": "I'm suspicious, test loyalty constantly",
            "Schizoid": "I prefer distance, intimacy feels intrusive",
            "Hysterical": "Intense, dramatic, quickly attached",
            "Borderline": "Intense love-hate swings, fear of abandonment",
            "Masochistic": "I stay in painful situations too long"
        }
    },
    {
        "id": 4,
        "category": "Anger Expression",
        "question": "When angry, you:",
        "options": {
            "Narcissistic": "Feel insulted, lash out or withdraw coldly",
            "Obsessive": "Suppress it, become more rigid",
            "Depressive": "Turn it inward, feel guilty",
            "Paranoid": "Blame others, see attacks everywhere",
            "Schizoid": "Disconnect emotionally",
            "Hysterical": "Express dramatically, then forget quickly",
            "Borderline": "Rage intensely, then feel shame",
            "Masochistic": "Swallow it, express passive-aggressively"
        }
    },
    {
        "id": 5,
        "category": "Identity Sense",
        "question": "How stable is your sense of self?",
        "options": {
            "Narcissistic": "Grandiose but fragile underneath",
            "Obsessive": "Stable, defined by principles/roles",
            "Depressive": "Stable but fundamentally flawed",
            "Paranoid": "Stable, misunderstood by hostile world",
            "Schizoid": "Stable but detached observer",
            "Hysterical": "Defined by relationships, dramatic",
            "Borderline": "Fragmented, changes with mood/person",
            "Masochistic": "Defined by suffering, victimhood"
        }
    },
    {
        "id": 6,
        "category": "Criticism Response",
        "question": "When someone criticizes you:",
        "options": {
            "Narcissistic": "I feel shattered or enraged, question their competence",
            "Obsessive": "I defend with logic, prove I was right",
            "Depressive": "I agree, always knew I wasn't good enough",
            "Paranoid": "I knew they were against me all along",
            "Schizoid": "I don't care, doesn't penetrate",
            "Hysterical": "I feel devastated, cry, need reassurance immediately",
            "Borderline": "I'm destroyed or you're evil—extreme reaction",
            "Masochistic": "I probably deserve it, apologize excessively"
        }
    },
    {
        "id": 7,
        "category": "Success/Achievement",
        "question": "When you succeed:",
        "options": {
            "Narcissistic": "I deserve it, should have more recognition",
            "Obsessive": "Relief, but worry about next task",
            "Depressive": "It was luck/others' help, feel guilty for pride",
            "Paranoid": "Finally proved them wrong, stay vigilant",
            "Schizoid": "Indifferent, just met requirements",
            "Hysterical": "Excitement, share dramatically with everyone",
            "Borderline": "Brief high, then empty or suspicious of praise",
            "Masochistic": "Uncomfortable, downplay or sabotage it"
        }
    },
    {
        "id": 8,
        "category": "Alone Time",
        "question": "Being alone for extended periods:",
        "options": {
            "Narcissistic": "Boring without audience, feel depleted",
            "Obsessive": "Productive time, organize and plan",
            "Depressive": "Painful but familiar, ruminate",
            "Paranoid": "Safer but lonely, prepare defenses",
            "Schizoid": "Prefer it, recharge, feel authentic",
            "Hysterical": "Unbearable, need stimulation/connection",
            "Borderline": "Terrifying emptiness, might self-harm",
            "Masochistic": "Deserved isolation, martyr myself"
        }
    },
    {
        "id": 9,
        "category": "Dependency",
        "question": "Someone needing you constantly:",
        "options": {
            "Narcissistic": "Flattering at first, then burdensome drain",
            "Obsessive": "Structured help is fine, clinginess irritates",
            "Depressive": "Finally my purpose, give endlessly",
            "Paranoid": "Suspicious of motives, what do they want?",
            "Schizoid": "Suffocating, need escape",
            "Hysterical": "Love it, feel important and connected",
            "Borderline": "Cling back intensely or push away in panic",
            "Masochistic": "Accept it, complain but stay trapped"
        }
    },
    {
        "id": 10,
        "category": "Moral Transgression",
        "question": "If you broke your own moral code:",
        "options": {
            "Narcissistic": "Justified exception, rules for others",
            "Obsessive": "Severe guilt, confess, overcompensate",
            "Depressive": "Proof I'm fundamentally bad, spiral",
            "Paranoid": "It was defensive, they drove me to it",
            "Schizoid": "Detached observation, intellectualize",
            "Hysterical": "Dramatic guilt, seek forgiveness theatrically",
            "Borderline": "Self-loathing, suicidal thoughts or dissociate",
            "Masochistic": "Expected failure, punish myself"
        }
    },
    {
        "id": 11,
        "category": "Fantasy Life",
        "question": "Your recurring daydreams involve:",
        "options": {
            "Narcissistic": "Triumph, admiration, revenge on detractors",
            "Obsessive": "Perfect scenarios, things going exactly right",
            "Depressive": "Being saved, or saving others sacrificially",
            "Paranoid": "Exposure of plots, vindication, survival",
            "Schizoid": "Elaborate inner worlds, no people needed",
            "Hysterical": "Romance, drama, being center of attention",
            "Borderline": "Merging/fusion, or violent destruction",
            "Masochistic": "Suffering nobly, eventual recognition of pain"
        }
    },
    {
        "id": 12,
        "category": "Body Experience",
        "question": "Your relationship with your body:",
        "options": {
            "Narcissistic": "Extension of image, must be perfect/impressive",
            "Obsessive": "Control it, discipline, function over pleasure",
            "Depressive": "Heavy, sluggish, burdensome",
            "Paranoid": "Vigilant to signs, body betrays/warns",
            "Schizoid": "Detached, mechanical housing",
            "Hysterical": "Dramatic symptoms, highly responsive",
            "Borderline": "Volatile—love/hate, self-harm, disconnect",
            "Masochistic": "Endure pain, deny pleasure"
        }
    },
    {
        "id": 13,
        "category": "Change/Transition",
        "question": "Major life changes make you feel:",
        "options": {
            "Narcissistic": "Opportunity to shine or threat to status",
            "Obsessive": "Anxious, need to plan every detail",
            "Depressive": "Loss-focused, mourn what's ending",
            "Paranoid": "Suspicious, threatened, danger lurking",
            "Schizoid": "Indifferent if autonomy preserved",
            "Hysterical": "Excited/terrified, overly dramatic",
            "Borderline": "Panicked, identity-shattering chaos",
            "Masochistic": "Resigned suffering, inevitable hardship"
        }
    },
    {
        "id": 14,
        "category": "Competition",
        "question": "In competitive situations:",
        "options": {
            "Narcissistic": "Must win, cheating if needed, crushing defeat intolerable",
            "Obsessive": "Play by rules perfectly, resent rule-breakers",
            "Depressive": "Don't deserve to win, give up easily",
            "Paranoid": "Everyone cheats, must stay vigilant",
            "Schizoid": "Disinterested, pointless exercise",
            "Hysterical": "Dramatic display, enjoy attention more than winning",
            "Borderline": "All-or-nothing intensity, rage if losing",
            "Masochistic": "Lose nobly, winning feels wrong"
        }
    },
    {
        "id": 15,
        "category": "Authority Figures",
        "question": "Your relationship with authority:",
        "options": {
            "Narcissistic": "Respect only if superior, compete otherwise",
            "Obsessive": "Respect structure, obey rules, secretly resent",
            "Depressive": "Defer completely, seek approval desperately",
            "Paranoid": "Distrust, see hidden agendas, rebellion",
            "Schizoid": "Minimal engagement, comply superficially",
            "Hysterical": "Seductive or defiant, emotionally reactive",
            "Borderline": "Idealize then devalue rapidly",
            "Masochistic": "Submit, secretly resent, provoke punishment"
        }
    },
    {
        "id": 16,
        "category": "Emptiness/Meaning",
        "question": "Sense of inner emptiness:",
        "options": {
            "Narcissistic": "When not admired, brief then refill externally",
            "Obsessive": "Rare, stay busy to avoid",
            "Depressive": "Chronic, life feels meaningless",
            "Paranoid": "Filled with vigilance, no room for emptiness",
            "Schizoid": "Comfortable void, prefer it",
            "Hysterical": "Terrifying, fill with drama/relationships",
            "Borderline": "Pervasive black hole, unbearable",
            "Masochistic": "Filled with suffering, gives purpose"
        }
    },
    {
        "id": 17,
        "category": "Jealousy/Envy",
        "question": "When others have what you want:",
        "options": {
            "Narcissistic": "Intense envy, devalue their achievement",
            "Obsessive": "Work harder, deserve it through effort",
            "Depressive": "They deserve it, I don't",
            "Paranoid": "They stole/cheated their way",
            "Schizoid": "Detached observation, don't really want it",
            "Hysterical": "Dramatic display of feeling left out",
            "Borderline": "Consuming rage or self-destruction",
            "Masochistic": "Expected, martyrdom continues"
        }
    },
    {
        "id": 18,
        "category": "Trust",
        "question": "You trust others:",
        "options": {
            "Narcissistic": "Only when they validate me consistently",
            "Obsessive": "When proven reliable through time",
            "Depressive": "Too much, get hurt repeatedly",
            "Paranoid": "Never fully, always testing",
            "Schizoid": "Don't need to, keep distance",
            "Hysterical": "Immediately and completely, then betrayed",
            "Borderline": "Desperately then not at all, oscillating",
            "Masochistic": "Despite betrayal, stay loyal to pain"
        }
    }
]


# ---------------------------------------------------------------------------
# Chatbot personality definitions
# ---------------------------------------------------------------------------

CHATBOT_PERSONALITIES = {
    "Psychoanalytic": {
        "prompt": (
            "You engage in a formal psychoanalytic mode — exploring unconscious patterns, defences, "
            "object relations, and early developmental themes. You interpret, make connections across "
            "time, and help the person gain insight into deeper psychological structures."
        ),
        "tts_tld": "com",
        "tts_slow": False,
    },
    "Supportive": {
        "prompt": (
            "You are warm, validating, and encouraging. You reflect the person's feelings back to them, "
            "normalise their experiences, and offer genuine emotional support. You focus on strengths "
            "and resilience alongside challenges. You are never clinical or cold."
        ),
        "tts_tld": "co.uk",
        "tts_slow": False,
    },
    "Socratic": {
        "prompt": (
            "You respond primarily with thoughtful, open-ended questions. You rarely give direct answers — "
            "instead you guide the person to discover insights themselves. Each response should contain at "
            "least one significant question that invites genuine self-examination."
        ),
        "tts_tld": "com",
        "tts_slow": True,
    },
    "Direct Coach": {
        "prompt": (
            "You are action-oriented and forward-focused. You acknowledge patterns briefly, then pivot to "
            "practical implications and behavioural change. You are concise, clear, and encouraging "
            "without being overly warm. You give specific, concrete suggestions."
        ),
        "tts_tld": "com.au",
        "tts_slow": False,
    },
    "Empathic Listener": {
        "prompt": (
            "You practise deep reflective listening. You reflect back what you hear — both the content "
            "and the emotional undertone — before gently exploring. You validate before you question. "
            "You make the person feel profoundly heard and understood above all else."
        ),
        "tts_tld": "co.uk",
        "tts_slow": False,
    },
    "Child Friendly (8–12)": {
        "prompt": (
            "You are talking with a child aged roughly 8 to 12. Use simple, everyday words and short sentences. "
            "Never use clinical or psychological jargon — if you need to explain a feeling or pattern, use a "
            "concrete analogy or a story they can picture (e.g. 'it's like when...'). Be warm, patient, and "
            "encouraging. Celebrate what they share. Ask one simple, friendly question at a time. "
            "If they seem confused, rephrase immediately. Make the conversation feel safe and fun, not like a test."
        ),
        "tts_tld": "com",
        "tts_slow": True,
    },
    "Teen (13–17)": {
        "prompt": (
            "You are talking with a teenager aged roughly 13 to 17. Be genuine and direct — teens spot "
            "condescension immediately, so never talk down to them or over-explain. Validate their feelings "
            "strongly and take their experiences seriously. Use plain, natural language (not overly formal, "
            "not trying too hard to sound young). Avoid heavy clinical language; if you use a concept, "
            "explain it briefly in real terms. Be curious about their perspective rather than prescriptive. "
            "Acknowledge that things can be confusing or contradictory and that's completely normal."
        ),
        "tts_tld": "com",
        "tts_slow": False,
    },
}

PERSONALITY_NAMES = list(CHATBOT_PERSONALITIES.keys())


# ---------------------------------------------------------------------------
# API
# ---------------------------------------------------------------------------

def call_claude_api(prompt=None, messages=None, system=None, max_tokens=2000):
    try:
        msg_list = [{"role": "user", "content": prompt}] if prompt is not None else (messages or [])
        payload = {
            "model": "claude-sonnet-4-20250514",
            "max_tokens": max_tokens,
            "messages": msg_list
        }
        if system:
            payload["system"] = system

        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "anthropic-version": "2023-06-01",
                "x-api-key": f"{os.getenv('ANTHROPIC_API_KEY')}"
            },
            json=payload
        )
        if response.status_code == 200:
            return response.json()['content'][0]['text']
        else:
            return f"API Error: {response.status_code} - {response.text}\n"
    except Exception as e:
        return f"Error calling API: {str(e)}\n"


def _build_context_section(extra_info="", memories=None):
    section = ""
    if extra_info and extra_info.strip():
        section += f"\n\n**Additional Information About This Person:**\n{extra_info.strip()}"
    if memories:
        mem_texts = [m['content'] for m in memories if m.get('content')]
        if mem_texts:
            section += "\n\n**Key Insights from Previous Conversations:**\n"
            for mem in mem_texts:
                section += f"- {mem}\n"
    return section


def analyze_narrative(responses_json, extra_info="", memories=None):
    context = _build_context_section(extra_info, memories)
    prompt = f"""You are a senior clinical psychologist specializing in psychoanalytic personality assessment. \
Write a formal, comprehensive personality assessment report in professional clinical prose. \
Use flowing paragraphs — not bullet points or numbered lists.

Assessment data:

{responses_json}{context}

Structure the report with these clearly headed prose sections:

**Executive Summary**
Primary personality organization, confidence level, and a one-sentence characterization of the core structure.

**Primary Personality Organization**
Dominant type, theoretical basis, how the responses support it, and level of personality organization \
(Neurotic / Borderline / Psychotic) with rationale.

**Secondary Features and Mixed Presentation**
Secondary patterns, how they interact with the primary type, contradictory or ambiguous patterns.

**Core Psychological Characteristics**
Defences, object relations, sense of self, emotional regulation style, relationship to anxiety — \
grounded in specific responses.

**Strengths and Adaptive Capacities**
Genuine psychological strengths and how they serve the individual.

**Vulnerabilities and Developmental Challenges**
Characteristic vulnerabilities, suffering patterns, and relational difficulties.

**Therapeutic Considerations**
Recommended approach, focus areas, and what this person may find challenging in therapy.

**Important Caveats**
Limitations of self-report assessment, contextual factors, and the broader clinical picture.

Use formal but accessible language. The report should be thorough."""

    return call_claude_api(prompt=prompt, max_tokens=3000)


def analyze_summary(responses_json, extra_info="", memories=None):
    context = _build_context_section(extra_info, memories)
    prompt = f"""You are a clinical psychologist. Provide a structured personality assessment summary.

{responses_json}{context}

Use clear markdown headings and concise bullet points:

## Primary Personality Type
- Dominant type, confidence (High/Medium/Low), one-line rationale
- Level of personality organization (Neurotic / Borderline / Psychotic)

## Secondary Features
- Secondary patterns present and how they interact with the primary type

## Pattern Analysis
- Most indicative assessment responses
- Contradictory or ambiguous patterns

## Key Characteristics
- 5–7 core traits evidenced by the responses

## Strengths
- 3–4 adaptive strengths

## Challenges
- 3–4 characteristic vulnerabilities

## Therapeutic Considerations
- Recommended approach and key focus areas

## Important Notes
- Caveats and contextual qualifications"""

    return call_claude_api(prompt=prompt, max_tokens=2000)


def run_full_analysis(responses_json, extra_info="", memories=None):
    narrative = analyze_narrative(responses_json, extra_info, memories)
    summary = analyze_summary(responses_json, extra_info, memories)
    return narrative, summary


def fallback_analysis(responses_json):
    data = json.loads(responses_json)
    type_counts = {}
    for response in data['responses']:
        ptype = response['selected_type']
        type_counts[ptype] = type_counts.get(ptype, 0) + 1

    dominant_type = max(type_counts.items(), key=lambda x: x[1])
    total = len(data['responses'])
    percentage = (dominant_type[1] / total) * 100

    text = f"""## Personality Assessment Results

**Assessment Date**: {data['timestamp']}

### Primary Personality Type: {dominant_type[0]}
**Confidence**: {'High' if percentage > 60 else 'Medium' if percentage > 40 else 'Low'}
**Indicator Strength**: {dominant_type[1]}/{total} responses ({percentage:.1f}%)

### Response Distribution:
"""
    for ptype, count in sorted(type_counts.items(), key=lambda x: x[1], reverse=True):
        pct = (count / total) * 100
        text += f"- **{ptype}**: {count} responses ({pct:.1f}%)\n"

    text += f"""
### Analysis Summary:
Based on {total} assessment questions, the dominant personality pattern is **{dominant_type[0]}**.

**Note**: This is a simplified analysis. Set ANTHROPIC_API_KEY in .env for full LLM analysis.

### Questions Answered:
"""
    for response in data['responses']:
        text += f"\n- **Q{response['question_id']}** ({response['category']}): {response['selected_type']}"
    return text


# ---------------------------------------------------------------------------
# UI helpers
# ---------------------------------------------------------------------------

def create_buttons_attributes(questions, input_responses=None):
    buttons = {}
    if questions:
        for question in questions:
            responses = list(question['options'].values())
            random.shuffle(responses)
            response = None
            if input_responses:
                response = next((r for r in input_responses if r["question_id"] == question['id']), None)
            buttons[question['id']] = {
                "choices": responses,
                "label": "Select your response:",
                "interactive": True,
                "value": response["selected_answer"] if response else None
            }
    return buttons


# ---------------------------------------------------------------------------
# Profile management
# ---------------------------------------------------------------------------

def create_profile(assessment_data, summary, narrative="", subject_name="",
                   extra_info="", memories=None, chat_history=None):
    return {
        "profile_version": "1.0",
        "profile_id": str(uuid.uuid4()),
        "subject_name": subject_name,
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),
        "last_chat_at": None,
        "assessment": assessment_data,
        "analysis": summary,
        "analysis_narrative": narrative,
        "analysis_summary": summary,
        "extra_info": extra_info,
        "memories": memories or [],
        "chat_history": chat_history or []
    }


def extract_memories(text):
    pattern = r'\[MEMORY:\s*(.*?)\]'
    memories = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
    cleaned = re.sub(pattern, '', text, flags=re.IGNORECASE | re.DOTALL).strip()
    return cleaned, [m.strip() for m in memories if m.strip()]


def format_memories_display(profile):
    if not profile:
        return "No memories saved yet."
    mems = profile.get('memories', [])
    if not mems:
        return "No memories saved yet."
    lines = []
    for m in mems:
        ts = m.get('timestamp', '')[:10]
        lines.append(f"**[{ts}]** {m.get('content', '')}")
    return "\n\n".join(lines)


def build_chat_system_prompt(profile, chattiness=3, personality="Psychoanalytic"):
    if chattiness <= 1:
        style = "Be extremely concise — one or two sentences maximum unless the question genuinely requires more."
    elif chattiness == 2:
        style = "Be concise and focused — a short paragraph. Add one observation only if it truly adds value."
    elif chattiness == 4:
        style = "Be warm, conversational, and thorough. Connect themes, ask a thoughtful follow-up question."
    elif chattiness >= 5:
        style = (
            "Be expansive, reflective, and deeply engaged. Offer rich insight, explore nuance, "
            "draw connections across the profile, and invite genuine dialogue with follow-up questions."
        )
    else:
        style = "Maintain a balanced tone — typically two to three paragraphs, substantive but not exhaustive."

    persona = CHATBOT_PERSONALITIES.get(personality, CHATBOT_PERSONALITIES["Psychoanalytic"])

    lines = [
        "You are a psychoanalytic consultant engaging with someone who has completed a formal personality assessment.",
        "",
        f"Therapeutic approach: {persona['prompt']}",
        "",
        f"Response length: {style}",
        "",
        "MEMORY INSTRUCTIONS — critical:",
        "Actively capture what the person reveals about themselves during conversation.",
        "When they share concrete information — specific behaviours, relationships, experiences,",
        "recurring patterns, significant events, or how they perceive themselves — save it immediately:",
        "[MEMORY: <brief factual observation about the person>]",
        "Include multiple [MEMORY: ...] tags as needed. Err on the side of saving too much.",
        "Only skip if the person shared nothing new in this exchange.",
        "",
        "## Current Profile",
        "",
    ]

    if profile.get('subject_name'):
        lines.append(f"**Name**: {profile['subject_name']}")

    lines.append(f"**Assessment Date**: {profile.get('assessment', {}).get('timestamp', 'Unknown')}")
    lines.append("")
    lines.append("## Personality Analysis (Summary)")
    lines.append(profile.get('analysis_summary', profile.get('analysis', 'No analysis available yet.')))

    if profile.get('extra_info', '').strip():
        lines.append("")
        lines.append("## Additional Information Provided")
        lines.append(profile['extra_info'])

    if profile.get('memories'):
        lines.append("")
        lines.append("## Saved Memories from Prior Conversations")
        for m in profile['memories']:
            ts = m.get('timestamp', '')[:10]
            lines.append(f"- [{ts}] {m.get('content', '')}")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Report view toggle
# ---------------------------------------------------------------------------

def switch_report_view(view_choice, profile):
    if not profile:
        return ""
    if view_choice == "Summary":
        return profile.get("analysis_summary", profile.get("analysis", "No summary available."))
    return profile.get("analysis_narrative", profile.get("analysis", "No full report available."))


# ---------------------------------------------------------------------------
# Assessment processing
# ---------------------------------------------------------------------------

def process_assessment(*args):
    radio_values = args[:len(questions)]
    extra_info = args[len(questions)] if len(args) > len(questions) else ""
    existing_profile = args[len(questions) + 1] if len(args) > len(questions) + 1 else {}

    responses = []
    for i, answer in enumerate(radio_values):
        if answer:
            q = questions[i]
            selected_type = None
            for ptype, text in q['options'].items():
                if text == answer:
                    selected_type = ptype
                    break
            responses.append({
                "question_id": q['id'],
                "category": q['category'],
                "question": q['question'],
                "selected_answer": answer,
                "selected_type": selected_type
            })

    if len(responses) < len(questions):
        return "⚠️ Please answer all questions before submitting.", "", {}, "Full Report"

    assessment_data = {
        "timestamp": datetime.now().isoformat(),
        "total_questions": len(questions),
        "responses": responses
    }
    json_output = json.dumps(assessment_data, indent=2)

    memories = existing_profile.get('memories', []) if existing_profile else []
    extra_info = extra_info or (existing_profile.get('extra_info', '') if existing_profile else '')

    narrative, summary = run_full_analysis(json_output, extra_info=extra_info, memories=memories)

    subject_name = existing_profile.get('subject_name', '') if existing_profile else ''
    profile = create_profile(assessment_data, summary, narrative=narrative,
                             subject_name=subject_name, extra_info=extra_info, memories=memories)

    return narrative, json_output, profile, "Full Report"


# ---------------------------------------------------------------------------
# Voice: speech-to-text and text-to-speech
# ---------------------------------------------------------------------------

def transcribe_speech(audio_path):
    """Convert recorded audio to text. Uses OpenAI Whisper if available, else Google STT."""
    if audio_path is None:
        return ""

    # Try OpenAI Whisper first (if both key and package are present)
    if os.getenv('OPENAI_API_KEY'):
        try:
            import openai
            client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
            with open(audio_path, 'rb') as f:
                transcript = client.audio.transcriptions.create(model="whisper-1", file=f)
            return transcript.text.strip()
        except ImportError:
            pass  # openai package not installed — fall through to Google STT
        except Exception as e:
            return f"⚠️ Transcription error: {str(e)}"

    # Fallback: Google STT via SpeechRecognition (free, no key needed)
    try:
        import speech_recognition as sr
        r = sr.Recognizer()
        with sr.AudioFile(audio_path) as source:
            audio_data = r.record(source)
        return r.recognize_google(audio_data).strip()
    except Exception as e:
        return f"⚠️ Transcription error: {str(e)}"


def speak_response(text, personality="Psychoanalytic"):
    """Convert text to speech and return path to audio file."""
    try:
        from gtts import gTTS

        # Strip markdown and memory tags for cleaner speech
        clean = re.sub(r'\[MEMORY:.*?\]', '', text, flags=re.IGNORECASE | re.DOTALL)
        clean = re.sub(r'[#*`_]', '', clean)
        clean = re.sub(r'\n+', ' ', clean).strip()
        if not clean:
            return None

        # Cap length to avoid very long TTS calls
        if len(clean) > 2500:
            clean = clean[:2500] + "..."

        cfg = CHATBOT_PERSONALITIES.get(personality, CHATBOT_PERSONALITIES["Psychoanalytic"])
        tts = gTTS(text=clean, lang='en', tld=cfg['tts_tld'], slow=cfg['tts_slow'])

        tmp = tempfile.NamedTemporaryFile(suffix='.mp3', delete=False)
        tts.save(tmp.name)
        tmp.close()
        return tmp.name
    except Exception:
        return None  # Fail silently — text response still visible


# ---------------------------------------------------------------------------
# Chat initialisation
# ---------------------------------------------------------------------------

def generate_chat_welcome(profile):
    name = profile.get('subject_name', '')
    last_chat = profile.get('last_chat_at')
    analysis_summary = profile.get('analysis_summary', profile.get('analysis', ''))
    memories = profile.get('memories', [])

    if last_chat:
        try:
            dt = datetime.fromisoformat(last_chat)
            time_context = f"Their last chat session was on {dt.strftime('%A, %B %d at %I:%M %p')}."
        except Exception:
            time_context = "They have chatted before."
    else:
        time_context = "This is their first chat session."

    memory_context = ""
    if memories:
        recent = memories[-5:]
        memory_context = "\n\nKey things remembered from previous conversations:\n" + \
                         "\n".join(f"- {m['content']}" for m in recent)

    name_line = f"Their name is {name}." if name else "Their name is not known."

    prompt = f"""You are a psychoanalytic consultant beginning a chat session with a client.

{name_line}
{time_context}

Their personality profile summary:
{analysis_summary[:1200]}
{memory_context}

Write a warm, personalised opening message. It should:
- Greet them {'by name' if name else 'warmly'}
- If this is a return visit, acknowledge the time since the last session naturally, mentioning the date
- If there are prior memories, reference a relevant theme briefly and naturally (not verbatim)
- Invite them into the conversation by naming something genuinely interesting from their profile
- Close with a single thoughtful, open-ended question specific to their personality type

Keep it to 2–3 paragraphs. Be warm and human, not clinical. Do not include [MEMORY: ...] tags."""

    return call_claude_api(prompt=prompt, max_tokens=500)


def initialize_chat(history, profile):
    """Generate welcome when Chat tab is opened, only if chat is empty."""
    if history:
        return history, profile

    if not profile or not profile.get('assessment'):
        return [{
            "role": "assistant",
            "content": (
                "👋 Welcome! Please complete the personality assessment first "
                "(Assessment tab), then come back here to chat about your profile."
            )
        }], profile

    welcome_msg = generate_chat_welcome(profile)

    profile = dict(profile)
    profile['last_chat_at'] = datetime.now().isoformat()
    profile['updated_at'] = datetime.now().isoformat()

    new_history = [{"role": "assistant", "content": welcome_msg}]
    profile['chat_history'] = new_history

    return new_history, profile


# ---------------------------------------------------------------------------
# Chat
# ---------------------------------------------------------------------------

def chat_with_profile(message, history, profile, chattiness=3,
                      personality="Psychoanalytic", voice_enabled=False):
    if not message or not message.strip():
        return history, "", profile, format_memories_display(profile), None

    if not profile or not profile.get('assessment'):
        new_history = list(history) + [
            {"role": "user", "content": message},
            {"role": "assistant", "content": (
                "⚠️ Please complete a personality assessment first (Assessment tab), "
                "then return here to chat about your profile."
            )}
        ]
        return new_history, "", profile, format_memories_display(profile), None

    system = build_chat_system_prompt(profile, chattiness=int(chattiness), personality=personality)
    messages = [{"role": h["role"], "content": h["content"]} for h in history]
    messages.append({"role": "user", "content": message})

    response = call_claude_api(messages=messages, system=system, max_tokens=1500)
    cleaned_response, new_memories = extract_memories(response)

    profile = dict(profile)
    if new_memories:
        existing = list(profile.get('memories', []))
        for mem in new_memories:
            existing.append({"timestamp": datetime.now().isoformat(), "content": mem})
        profile['memories'] = existing
        profile['updated_at'] = datetime.now().isoformat()

    new_history = list(history) + [
        {"role": "user", "content": message},
        {"role": "assistant", "content": cleaned_response}
    ]
    profile['chat_history'] = new_history

    voice_audio = speak_response(cleaned_response, personality) if voice_enabled else None

    return new_history, "", profile, format_memories_display(profile), voice_audio


def handle_voice_input(audio_path, history, profile, chattiness, personality, voice_enabled):
    """Transcribe microphone input, send to chat, return response (always with voice)."""
    if audio_path is None:
        return history, None, profile, format_memories_display(profile)

    text = transcribe_speech(audio_path)

    if not text or text.startswith("⚠️"):
        error_msg = text or "⚠️ Could not transcribe audio. Please try again."
        new_history = list(history) + [{"role": "assistant", "content": error_msg}]
        return new_history, None, profile, format_memories_display(profile)

    # Voice input always produces voice output
    new_history, _, profile, memories_disp, voice_audio = chat_with_profile(
        text, history, profile, chattiness, personality, voice_enabled=True
    )
    return new_history, voice_audio, profile, memories_disp


def reanalyze_from_chat(history, profile):
    if not profile or not profile.get('assessment'):
        msg = "⚠️ No assessment data found. Please complete a personality assessment first."
        return list(history) + [{"role": "assistant", "content": msg}], profile

    json_data = json.dumps(profile['assessment'], indent=2)
    narrative, summary = run_full_analysis(
        json_data,
        extra_info=profile.get('extra_info', ''),
        memories=profile.get('memories', [])
    )

    profile = dict(profile)
    profile['analysis_narrative'] = narrative
    profile['analysis_summary'] = summary
    profile['analysis'] = summary
    profile['updated_at'] = datetime.now().isoformat()

    chat_msg = (
        "I've re-analyzed your personality profile, incorporating all additional information "
        "and insights from our conversations. Here's the updated full report:\n\n" + narrative
    )
    new_history = list(history) + [{"role": "assistant", "content": chat_msg}]
    profile['chat_history'] = new_history
    return new_history, profile


def clear_memories(profile):
    if not profile:
        return profile, "No memories saved yet."
    profile = dict(profile)
    profile['memories'] = []
    profile['updated_at'] = datetime.now().isoformat()
    return profile, "No memories saved yet."


# ---------------------------------------------------------------------------
# Profile save / load
# ---------------------------------------------------------------------------

def _profile_to_saveable(profile):
    """Remove narrative (not stored) and ensure summary is canonical analysis."""
    p = dict(profile)
    p.pop('analysis_narrative', None)
    p['analysis'] = p.get('analysis_summary', p.get('analysis', ''))
    return p


def save_profile_as_json(profile, subject_name, filename_input):
    if not profile or not profile.get('assessment'):
        return None, "⚠️ No profile to save. Complete an assessment first."

    profile = dict(profile)
    if subject_name and subject_name.strip():
        profile['subject_name'] = subject_name.strip()
    profile['updated_at'] = datetime.now().isoformat()

    saveable = _profile_to_saveable(profile)

    tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8')
    json.dump(saveable, tmp, indent=2)
    tmp.close()

    if filename_input and filename_input.strip():
        display_name = filename_input.strip()
        if not display_name.endswith('.json'):
            display_name += '.json'
    else:
        name = profile.get('subject_name', '') or 'profile'
        safe = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_')
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        display_name = f"personality_profile_{safe}_{ts}.json" if safe else f"personality_profile_{ts}.json"

    return tmp.name, f"✅ Profile ready: **{display_name}**"


def load_profile_file(filepath):
    num_q = len(questions)
    empty = [None] * num_q

    try:
        if filepath is None:
            return {}, "⚠️ Please upload a file.", "", "", "", "", [], *empty

        original_filename = Path(filepath).name

        with open(filepath, encoding='utf-8') as f:
            data = json.load(f)

        if 'profile_version' in data:
            profile = data
            assessment_data = data.get('assessment', {})
            analysis_summary = data.get('analysis_summary', data.get('analysis', ''))
            analysis_narrative = data.get('analysis_narrative', '')
            extra_info = data.get('extra_info', '')
            chat_history = data.get('chat_history', [])
            profile['analysis_summary'] = analysis_summary
            profile['analysis_narrative'] = analysis_narrative
        elif 'responses' in data:
            assessment_data = data
            analysis_summary = ''
            analysis_narrative = ''
            extra_info = ''
            chat_history = []
            profile = create_profile(assessment_data, analysis_summary)
        else:
            return {}, "⚠️ Invalid file format.", "", "", "", "", [], *empty

        json_output = json.dumps(assessment_data, indent=2)
        buttons = create_buttons_attributes(questions, assessment_data.get('responses', []))
        button_values = [buttons[q['id']]['value'] for q in questions]

        subject_name = profile.get('subject_name', '')
        status = f"✅ Profile loaded: **{subject_name or 'unnamed'}**"

        return (profile, status, analysis_summary, json_output, extra_info,
                original_filename, chat_history, *button_values)

    except json.JSONDecodeError:
        return {}, "⚠️ Invalid JSON file.", "", "", "", "", [], *empty
    except Exception as e:
        return {}, f"⚠️ Error loading file: {str(e)}", "", "", "", "", [], *empty


# ---------------------------------------------------------------------------
# Export functions
# ---------------------------------------------------------------------------

def export_to_word(analysis_text, json_data, custom_filename=None, save_location=None):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading('Psychoanalytic Personality Assessment Results', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if json_data:
            try:
                data = json.loads(json_data)
                p = doc.add_paragraph(f'Assessment Date: {data.get("timestamp", "")}')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

        doc.add_paragraph()
        doc.add_heading('Analysis', level=1)

        for line in analysis_text.split('\n'):
            if line.strip():
                if line.startswith('###'):
                    doc.add_heading(line.replace('#', '').strip(), level=2)
                elif line.startswith('##'):
                    doc.add_heading(line.replace('#', '').strip(), level=1)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(line.replace('**', '')).bold = True
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)

        if json_data:
            try:
                data = json.loads(json_data)
                doc.add_page_break()
                doc.add_heading('Detailed Response Data', level=1)
                for resp in data['responses']:
                    p = doc.add_paragraph()
                    p.add_run(f"Question {resp['question_id']}: ").bold = True
                    p.add_run(f"{resp['category']}\n")
                    p.add_run("Selected Type: ").bold = True
                    p.add_run(f"{resp['selected_type']}\n")
                    if 'selected_answer' in resp:
                        p.add_run("Answer: ").bold = True
                        p.add_run(f"{resp['selected_answer']}\n")
                    doc.add_paragraph()
            except Exception:
                pass

        filename = (custom_filename.strip() if custom_filename and custom_filename.strip()
                    else f"assessment_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        if not filename.endswith('.docx'):
            filename += '.docx'

        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(tmp.name)
        tmp.close()
        return tmp.name

    except ImportError:
        os.system("pip install python-docx --break-system-packages -q")
        return export_to_word(analysis_text, json_data, custom_filename, save_location)
    except Exception as e:
        return f"Error creating Word document: {str(e)}"


def export_to_markdown(analysis_text, json_data, custom_filename=None, save_location=None):
    try:
        md = "# Psychoanalytic Personality Assessment Results\n\n"
        if json_data:
            try:
                data = json.loads(json_data)
                md += f"**Assessment Date**: {data.get('timestamp', '')}\n\n"
            except Exception:
                pass

        md += "---\n\n## Analysis\n\n" + analysis_text + "\n\n"

        if json_data:
            try:
                data = json.loads(json_data)
                md += "---\n\n## Detailed Response Data\n\n"
                for resp in data['responses']:
                    md += f"### Question {resp['question_id']}: {resp['category']}\n\n"
                    md += f"**Selected Type**: {resp['selected_type']}\n\n"
                    if 'selected_answer' in resp:
                        md += f"**Answer**: {resp['selected_answer']}\n\n"
                    md += "---\n\n"
            except Exception:
                pass

        filename = (custom_filename.strip() if custom_filename and custom_filename.strip()
                    else f"assessment_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md")
        if not filename.endswith('.md'):
            filename += '.md'

        tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8')
        tmp.write(md)
        tmp.close()
        return tmp.name
    except Exception as e:
        return f"Error creating Markdown file: {str(e)}"


def generate_suggested_filename(extension='docx'):
    return f"assessment_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{extension}"


def get_default_save_location():
    home = str(Path.home())
    for candidate in [os.path.join(home, "Documents"), os.path.join(home, "Desktop"),
                      os.path.join(home, "Downloads")]:
        if os.path.exists(candidate):
            return candidate
    return home


def get_save_location_suggestions():
    home = str(Path.home())
    suggestions = []
    for p in [os.path.join(home, "Documents"), os.path.join(home, "Downloads"),
              os.path.join(home, "Desktop")]:
        if os.path.exists(p):
            suggestions.append(p)
    suggestions.append(home)
    return suggestions


# ---------------------------------------------------------------------------
# Gradio UI
# ---------------------------------------------------------------------------

with gr.Blocks(title="Psychoanalytic Personality Assessment") as demo:

    profile_state = gr.State({})

    gr.Markdown("""
# 🧠 Psychoanalytic Personality Assessment

This comprehensive assessment uses psychoanalytic theory to identify personality patterns.
Answer all 18 questions honestly based on how you typically think, feel, and behave.
""")

    with gr.Tabs():

        # ── Tab 1: Assessment ────────────────────────────────────────────────
        with gr.Tab("📝 Assessment"):
            buttons = create_buttons_attributes(questions)
            radio_buttons = []

            with gr.Tabs():
                for tab_num in range(3):
                    start_idx = tab_num * 6
                    end_idx = min((tab_num + 1) * 6, len(questions))
                    with gr.Tab(f"Questions {start_idx + 1}–{end_idx}"):
                        for q in questions[start_idx:end_idx]:
                            with gr.Group():
                                gr.Markdown(f"### Question {q['id']}: {q['category']}")
                                gr.Markdown(f"**{q['question']}**")
                                radio = gr.Radio(
                                    choices=buttons[q['id']]['choices'],
                                    label=buttons[q['id']]['label'],
                                    interactive=True,
                                    value=None
                                )
                                radio_buttons.append(radio)

            gr.Markdown("---")
            gr.Markdown("## 📝 Additional Information")
            gr.Markdown(
                "Optionally provide extra context before running the analysis. "
                "Existing profile context and conversation memories are included automatically."
            )
            extra_info_input = gr.Textbox(
                label="Additional Information",
                placeholder="Background, observed behaviours, history, or any other relevant context…",
                lines=4
            )

            gr.Markdown("---")
            with gr.Row():
                submit_btn = gr.Button("📊 Analyze My Personality", variant="primary", size="lg")
                clear_btn = gr.Button("🔄 Clear All Responses", size="lg")

            gr.Markdown("---")
            gr.Markdown("## 📋 Assessment Results")

            with gr.Row():
                report_view_toggle = gr.Radio(
                    choices=["Full Report", "Summary"],
                    value="Full Report",
                    label="Report View",
                    interactive=True
                )

            with gr.Row():
                with gr.Column(scale=2):
                    analysis_output = gr.Markdown(label="Personality Analysis")
                with gr.Column(scale=1):
                    json_output = gr.Code(label="Response Data (JSON)", language="json")

        # ── Tab 2: Profile ───────────────────────────────────────────────────
        with gr.Tab("👤 Profile"):
            gr.Markdown("## Save & Load Profile")
            gr.Markdown(
                "A **profile** bundles assessment responses, the summary analysis, additional information, "
                "chat memories and conversation history. The narrative report is generated fresh each session "
                "and is not stored."
            )

            gr.Markdown("### 💾 Save Current Profile")
            with gr.Row():
                subject_name_input = gr.Textbox(label="Subject Name (optional)",
                                                placeholder="Enter a name for this profile…")
            with gr.Row():
                profile_filename_input = gr.Textbox(
                    label="Save As (filename)",
                    placeholder="personality_profile.json",
                    info="Uploading a profile pre-fills this with the source filename."
                )
            with gr.Row():
                save_profile_btn = gr.Button("💾 Save Profile as JSON", variant="primary", size="lg")

            profile_save_status = gr.Markdown("")
            profile_file_output = gr.File(label="Download Profile JSON")

            gr.Markdown("---")
            gr.Markdown("### 📂 Load an Existing Profile")
            gr.Markdown(
                "Upload a previously saved profile JSON to restore the assessment, summary analysis, "
                "additional information, memories, and chat history."
            )
            profile_file_upload = gr.File(label="Upload Profile JSON", file_types=[".json"])
            profile_load_status = gr.Markdown("")

        # ── Tab 3: Chat ──────────────────────────────────────────────────────
        with gr.Tab("💬 Chat") as chat_tab:
            gr.Markdown("## Chat with Your Profile")
            gr.Markdown(
                "The AI has full access to your personality profile and remembers what you share. "
                "Speak or type — voice input always gets a spoken response."
            )

            with gr.Row():
                # ── Left: conversation area ──────────────────────────────────
                with gr.Column(scale=3):
                    chatbot = gr.Chatbot(
                        label="Conversation",
                        height=460,
                        render_markdown=True
                    )

                    gr.Markdown(
                        "**🎤 Voice input** — click Record, speak, then click Stop. "
                        "Your message is transcribed and sent automatically. "
                        "If the microphone is unavailable you can upload a recorded audio file instead. "
                        "*(Grant microphone access in your browser when prompted.)*"
                    )
                    with gr.Row():
                        voice_input = gr.Audio(
                            sources=["microphone", "upload"],
                            type="filepath",
                            label="🎤 Voice Input (record or upload)"
                        )
                        voice_output = gr.Audio(
                            label="🔊 AI Voice Response",
                            autoplay=True,
                            interactive=False,
                            sources=[]
                        )

                    chat_input = gr.Textbox(
                        label="Type a message  (Enter to send)",
                        placeholder="Ask about your personality, explore a pattern, share more context…",
                        lines=1,
                        max_lines=6
                    )

                    with gr.Row():
                        chat_clear_btn = gr.Button("Clear Chat")
                        chat_reanalyze_btn = gr.Button(
                            "🔄 Re-analyze with Chat Context", variant="secondary"
                        )

                # ── Right: controls ─────────────────────────────────────────
                with gr.Column(scale=1):
                    chatbot_personality = gr.Dropdown(
                        choices=PERSONALITY_NAMES,
                        value="Psychoanalytic",
                        label="Chatbot Personality",
                        info="Changes the AI's therapeutic style and voice accent."
                    )
                    chattiness_slider = gr.Slider(
                        minimum=1, maximum=5, value=3, step=1,
                        label="Response Length",
                        info="1 = Very concise  ·  3 = Balanced  ·  5 = Expansive"
                    )
                    voice_enabled_checkbox = gr.Checkbox(
                        label="🔊 Speak text replies",
                        value=False,
                        info="Enable to hear spoken responses when typing (voice input always speaks)."
                    )
                    gr.Markdown("---")
                    gr.Markdown("### 🧠 Saved Memories")
                    gr.Markdown(
                        "Meaningful insights the AI saves from your conversation "
                        "— included in future analyses."
                    )
                    memories_display = gr.Markdown("No memories saved yet.")
                    clear_memories_btn = gr.Button("🗑️ Clear Memories", size="sm")

        # ── Tab 4: Save Results ──────────────────────────────────────────────
        with gr.Tab("📤 Save Results"):
            gr.Markdown("### 💾 Export Analysis")

            with gr.Row():
                save_location_input = gr.Textbox(
                    label="📁 Save Location",
                    value=get_default_save_location(),
                    info=f"Default: {get_default_save_location()}"
                )

            with gr.Row():
                with gr.Column():
                    gr.Markdown("**Common Locations:**")
                    for loc in get_save_location_suggestions()[:4]:
                        gr.Markdown(f"• `{loc}`")

            with gr.Row():
                with gr.Column():
                    word_filename_input = gr.Textbox(
                        label="Word Filename",
                        value=generate_suggested_filename('docx'),
                        info="Extension .docx added automatically."
                    )
                    export_word_btn = gr.Button("📄 Export to Word (.docx)", size="lg")
                with gr.Column():
                    md_filename_input = gr.Textbox(
                        label="Markdown Filename",
                        value=generate_suggested_filename('md'),
                        info="Extension .md added automatically."
                    )
                    export_md_btn = gr.Button("📝 Export to Markdown (.md)", size="lg")

            gr.Markdown("**💡 Tip**: Files are downloaded via your browser.")

            with gr.Row():
                word_file_output = gr.File(label="Word Document")
                md_file_output = gr.File(label="Markdown Document")

    gr.Markdown("""
---
### 📌 Important Notes
- This assessment is for educational and self-reflection purposes only
- Results should not replace professional psychological evaluation
- All personality types have strengths and challenges; most people show mixed patterns

### 🔧 API & Voice Configuration
- Set `ANTHROPIC_API_KEY` in `.env` for LLM analysis (required)
- Voice input uses OpenAI Whisper if `OPENAI_API_KEY` is set, otherwise Google STT (free)
- Voice output requires `gTTS`: `pip install gTTS`
- STT fallback requires `SpeechRecognition`: `pip install SpeechRecognition`
""")

    # ── Event handlers ───────────────────────────────────────────────────────

    # Assessment
    submit_btn.click(
        fn=process_assessment,
        inputs=radio_buttons + [extra_info_input, profile_state],
        outputs=[analysis_output, json_output, profile_state, report_view_toggle]
    )

    clear_btn.click(
        fn=lambda: ([None] * len(questions) + ["", "", {}, "Full Report",
                    get_default_save_location(), generate_suggested_filename('docx'),
                    generate_suggested_filename('md'), None, None]),
        inputs=None,
        outputs=radio_buttons + [analysis_output, json_output, profile_state, report_view_toggle,
                                  save_location_input, word_filename_input, md_filename_input,
                                  word_file_output, md_file_output]
    )

    report_view_toggle.change(
        fn=switch_report_view,
        inputs=[report_view_toggle, profile_state],
        outputs=[analysis_output]
    )

    # Profile tab
    save_profile_btn.click(
        fn=save_profile_as_json,
        inputs=[profile_state, subject_name_input, profile_filename_input],
        outputs=[profile_file_output, profile_save_status]
    )

    profile_file_upload.change(
        fn=load_profile_file,
        inputs=[profile_file_upload],
        outputs=[profile_state, profile_load_status, analysis_output, json_output,
                 extra_info_input, profile_filename_input, chatbot, *radio_buttons]
    )

    # Chat tab — welcome on tab open
    chat_tab.select(
        fn=initialize_chat,
        inputs=[chatbot, profile_state],
        outputs=[chatbot, profile_state]
    )

    # Text chat — Enter sends
    chat_input.submit(
        fn=chat_with_profile,
        inputs=[chat_input, chatbot, profile_state, chattiness_slider,
                chatbot_personality, voice_enabled_checkbox],
        outputs=[chatbot, chat_input, profile_state, memories_display, voice_output]
    )

    # Voice chat — recording finished triggers transcribe + respond
    voice_input.change(
        fn=handle_voice_input,
        inputs=[voice_input, chatbot, profile_state, chattiness_slider,
                chatbot_personality, voice_enabled_checkbox],
        outputs=[chatbot, voice_output, profile_state, memories_display]
    )

    chat_clear_btn.click(
        fn=lambda profile: ([], None, profile),
        inputs=[profile_state],
        outputs=[chatbot, voice_output, profile_state]
    )

    chat_reanalyze_btn.click(
        fn=reanalyze_from_chat,
        inputs=[chatbot, profile_state],
        outputs=[chatbot, profile_state]
    )

    clear_memories_btn.click(
        fn=clear_memories,
        inputs=[profile_state],
        outputs=[profile_state, memories_display]
    )

    # Save Results tab
    export_word_btn.click(
        fn=export_to_word,
        inputs=[analysis_output, json_output, word_filename_input, save_location_input],
        outputs=[word_file_output]
    )

    export_md_btn.click(
        fn=export_to_markdown,
        inputs=[analysis_output, json_output, md_filename_input, save_location_input],
        outputs=[md_file_output]
    )


if __name__ == "__main__":
    load_dotenv()
    demo.launch()
